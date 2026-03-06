import os
import re
import json
import logging
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as pdf_extract
from docx import Document
from groq import Groq

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"pdf", "txt", "doc", "docx"}
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Get API key from environment
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
MODEL_NAME = "llama-3.3-70b-versatile"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    if not text:
        return ""
    text = "\n".join([line.strip() for line in text.splitlines() if line.strip()])
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file using python-docx"""
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    except Exception as e:
        logging.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_file(file_path):
    """Extract text from PDF, DOCX, or TXT files"""
    ext = file_path.rsplit(".", 1)[1].lower()
    
    try:
        if ext == "pdf":
            text = pdf_extract(file_path)
            return clean_text(text) if text.strip() else ""
            
        elif ext in ["doc", "docx"]:
            # Use python-docx directly (no LibreOffice needed!)
            text = extract_text_from_docx(file_path)
            return clean_text(text) if text.strip() else ""
            
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return clean_text(f.read().strip())
                
    except Exception as e:
        logging.error(f"Text extraction error: {e}")
    
    return ""

def parse_resume_with_ai(resume_text):
    """Parse resume using Groq AI"""
    
    if not GROQ_API_KEY:
        return {"error": "GROQ_API_KEY not configured"}
    
    prompt = f"""Extract information from this resume as JSON.

Resume:
{resume_text[:12000]}

Return ONLY valid JSON with these fields:
{{
  "personal_info": {{
    "first_name": null,
    "last_name": null,
    "email": null,
    "phone": null,
    "location": null,
    "linkedin": null,
    "github": null
  }},
  "education": [
    {{
      "degree": null,
      "field_of_study": null,
      "institution": null,
      "start_date": null,
      "end_date": null,
      "grade": null
    }}
  ],
  "experience": [
    {{
      "job_title": null,
      "company": null,
      "start_date": null,
      "end_date": null,
      "responsibilities": []
    }}
  ],
  "projects": [
    {{
      "name": null,
      "description": null,
      "technologies": []
    }}
  ],
  "skills": {{
    "technical_skills": [],
    "programming_languages": [],
    "frameworks": [],
    "tools": []
  }},
  "certifications": [],
  "hackathons": []
}}"""

    try:
        client = Groq(api_key=GROQ_API_KEY)
        
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": (
    "You are a resume parser. Follow these rules strictly:\n"
    "- Extract ALL job experiences from the resume, not just the first one\n"
    "- Extract ALL education entries\n"
    "- Extract ALL projects listed\n"
    "- For each experience entry, only include responsibilities specific to that job\n"
    "- Use the exact dates mentioned in the resume for start_date and end_date\n"
    "- If a field is genuinely missing, use null\n"
    "- Return only valid JSON with no explanation, preamble, or markdown formatting"
)},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=3000
        )
        
        content = response.choices[0].message.content
        content = re.sub(r'^```json\s*', '', content)
        content = re.sub(r'^```\s*', '', content)
        content = re.sub(r'\s*```$', '', content)
        content = content.strip()
        
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            content = json_match.group()
        
        parsed = json.loads(content)
        return parsed
        
    except Exception as e:
        logging.error(f"API error: {e}")
        return {"error": str(e)}

@app.route("/", methods=["GET"])
def index():
    return jsonify({
        "service": "Resume Parser API",
        "version": "1.0.0",
        "status": "active",
        "endpoints": {
            "/parse": {
                "method": "POST",
                "description": "Upload and parse resume",
                "supported_formats": ["pdf", "txt", "doc", "docx"]
            },
            "/health": {
                "method": "GET",
                "description": "Health check"
            }
        }
    }), 200

@app.route("/health", methods=["GET"])
def health():
    return jsonify({
        "status": "healthy",
        "groq_api_configured": bool(GROQ_API_KEY),
        "model": MODEL_NAME
    }), 200

@app.route("/parse", methods=["POST"])
def parse_resume():
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        
        if file.filename == "":
            return jsonify({"error": "No file selected"}), 400

        if not allowed_file(file.filename):
            return jsonify({
                "error": "Invalid file type",
                "allowed_types": list(ALLOWED_EXTENSIONS)
            }), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(file_path)
        logging.info(f"File saved: {file_path}")

        resume_text = extract_text_from_file(file_path)
        
        if not resume_text:
            return jsonify({"error": "Failed to extract text from resume"}), 500

        logging.info(f"Extracted {len(resume_text)} characters")

        parsed_data = parse_resume_with_ai(resume_text)

        # Clean up
        try:
            os.remove(file_path)
        except:
            pass

        return jsonify({
            "status": "success",
            "data": parsed_data,
            "metadata": {
                "filename": file.filename,
                "text_length": len(resume_text)
            }
        }), 200

    except Exception as e:
        logging.error(f"Error: {e}")
        return jsonify({
            "error": "Internal server error",
            "message": str(e)
        }), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
